"""
Meeting Report Generator - Creates new .docx report from previous + updates.

Strategy: Copy previous report .docx, then modify in-place.
This preserves all formatting, styles, and layout.

Key operations:
1. Update metadata (meeting number, date, distribution date)
2. Update attendance
3. Un-bold all previously bold text (demote from "latest" to "previous")
4. Append new paragraphs in bold (new meeting content)
5. Add new rows to section tables for new points
6. Update planning and info exchange tables
"""

import copy
import re
import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def copy_report(source_path, dest_path):
    """Copy the previous report as base for the new one."""
    shutil.copy2(source_path, dest_path)
    return Document(dest_path)


def _find_metadata_table(doc):
    """Find table 0 (metadata/header table)."""
    return doc.tables[0] if doc.tables else None


def _update_cell_text_preserve_format(cell, old_text_pattern, new_text):
    """Update text in a cell while preserving formatting.

    Finds runs matching old_text_pattern and replaces with new_text.
    For multi-character fields stored as individual character runs,
    clears all matching runs and writes new_text into the first one.
    """
    for para in cell.paragraphs:
        # Collect all run texts to find the pattern
        full_text = ''.join(r.text or '' for r in para.runs)
        match = re.search(old_text_pattern, full_text)
        if not match:
            continue

        start_pos = match.start()
        end_pos = match.end()

        # Find which runs correspond to the matched region
        pos = 0
        first_run_idx = None
        last_run_idx = None
        for ri, run in enumerate(para.runs):
            run_len = len(run.text or '')
            if pos + run_len > start_pos and first_run_idx is None:
                first_run_idx = ri
            if pos + run_len >= end_pos:
                last_run_idx = ri
                break
            pos += run_len

        if first_run_idx is not None and last_run_idx is not None:
            # Put new text in first matching run, clear the rest
            para.runs[first_run_idx].text = new_text
            for ri in range(first_run_idx + 1, last_run_idx + 1):
                para.runs[ri].text = ''
            return True

    return False


def update_metadata(doc, new_meeting_number, new_date, new_distribution_date=None):
    """Update the header table with new meeting number and date.

    Args:
        doc: Document object
        new_meeting_number: int, the new meeting number
        new_date: str, date in DD/MM/YYYY format
        new_distribution_date: str, optional distribution date
    """
    table = _find_metadata_table(doc)
    if not table:
        return

    # Row 1: Meeting number and date
    row1 = table.rows[1]

    # Update meeting number in cell 0 ("Minutes of meeting n° XX")
    # The number is stored across individual character runs
    cell0 = row1.cells[0]
    for para in cell0.paragraphs:
        full = ''.join(r.text or '' for r in para.runs)
        num_match = re.search(r'(\d+)$', full.strip())
        if num_match:
            old_num = num_match.group(1)
            _update_cell_text_preserve_format(
                cell0, re.escape(old_num) + r'$', str(new_meeting_number)
            )
            break

    # Update meeting number in the standalone number cell (cell 4 area)
    # This cell just contains the number
    for cell in row1.cells[1:]:
        cell_text = cell.text.strip()
        if cell_text.isdigit():
            for para in cell.paragraphs:
                if para.runs:
                    para.runs[0].text = str(new_meeting_number)
                    for r in para.runs[1:]:
                        r.text = ''
            break

    # Update date in cell that contains DD/MM/YYYY
    # Use search instead of match to handle "Du DD/MM/YYYY" (CORUM format)
    for cell in row1.cells:
        cell_text = cell.text.strip()
        if re.search(r'\d{2}/\d{2}/\d{4}', cell_text):
            _update_cell_text_preserve_format(
                cell, r'\d{2}/\d{2}/\d{4}', new_date
            )
            break

    # Row 2: Location (usually doesn't change, skip)

    # Row 3: Distribution date
    if new_distribution_date:
        row3 = table.rows[3]
        cell0 = row3.cells[0]
        _update_cell_text_preserve_format(
            cell0,
            r'\d{2}/\d{2}/\d{4}',
            new_distribution_date
        )


def update_next_meeting(doc, next_meeting_text, next_meeting_time=None):
    """Update the next meeting paragraph in the document.

    If next_meeting_text looks like a date (DD/MM/YYYY), only the date portion
    of the existing paragraph is replaced, preserving the original syntax.
    If next_meeting_time is provided (e.g. "11:00"), the time is also replaced.
    Otherwise the full paragraph text is replaced.
    """
    date_only = bool(re.match(r'^\d{2}/\d{2}/\d{4}$', next_meeting_text.strip()))

    found_label = False
    for p in doc.paragraphs:
        text = p.text.strip()
        if re.search(r'Next meeting|Prochaine r[eé]union', text, re.IGNORECASE):
            found_label = True
            continue
        if found_label and text:
            if date_only:
                # Replace only the date in the existing text, keep everything else
                full_text = ''.join(r.text or '' for r in p.runs)
                new_full = re.sub(r'\d{2}/\d{2}/\d{4}', next_meeting_text.strip(), full_text)
                # Also replace time if provided (matches "11:00", "11 :00", "11\xa0:00")
                if next_meeting_time:
                    new_full = re.sub(
                        r'(\d{1,2})[\s\xa0]*[:\.][\s\xa0]*(\d{2})',
                        next_meeting_time.strip(),
                        new_full,
                        count=1,
                    )
                if p.runs:
                    p.runs[0].text = new_full
                    for r in p.runs[1:]:
                        r.text = ''
            else:
                # Full replacement
                if p.runs:
                    p.runs[0].text = next_meeting_text
                    for r in p.runs[1:]:
                        r.text = ''
            return True
    return False


def unbold_all_content(doc):
    """Remove bold from all text in subject tables.

    This "demotes" the latest meeting content to normal weight,
    making room for the new meeting's content to be bold.
    """
    for table in doc.tables:
        # Only process subject tables (detected by header content)
        if len(table.rows) < 2:
            continue
        header_text = ' '.join(c.text.lower() for c in table.rows[0].cells)
        is_subject = ('subject' in header_text or 'sujet' in header_text
                      or 'n°' in header_text or re.search(r'd\d+\s*[-–]', header_text))
        if not is_subject:
            # Also unbold in planning table
            is_planning = 'planning' in header_text
            if not is_planning:
                continue

        for row in table.rows[1:]:
            for cell in row.cells:
                _unbold_cell(cell)


def _unbold_cell(cell):
    """Remove bold formatting from all runs in a cell."""
    for para in cell.paragraphs:
        for run in para.runs:
            if run.bold:
                run.bold = False


def add_paragraphs_to_cell(cell, paragraphs, bold=True):
    """Add new paragraphs to a cell.

    Args:
        cell: docx cell object
        paragraphs: list of text strings to add
        bold: whether new text should be bold
    """
    for text in paragraphs:
        new_para = cell.add_paragraph()
        run = new_para.add_run(text)
        run.bold = bold

        # Copy font properties from last existing paragraph
        existing_paras = cell.paragraphs
        if len(existing_paras) > 1:
            ref_para = existing_paras[-2]  # the one before our new one
            if ref_para.runs:
                ref_run = ref_para.runs[0]
                if ref_run.font.size:
                    run.font.size = ref_run.font.size
                if ref_run.font.name:
                    run.font.name = ref_run.font.name


def update_existing_point(table, point_number, updates, new_meeting_number):
    """Update an existing point in a subject table.

    Args:
        table: docx table containing the point
        point_number: str, e.g. "07.02"
        updates: dict with:
            - subject_lines: list of new text lines to append
            - for_whom: str, who is responsible
            - due: str, due date or status
        new_meeting_number: int, for the meeting header line
    """
    for row in table.rows[1:]:
        # Check actual XML cells for point number
        tr = row._tr
        tcs = tr.findall(qn('w:tc'))
        if not tcs:
            continue

        first_cell_text = _get_tc_text_simple(tcs[0]).strip()
        if first_cell_text != point_number:
            continue

        # Found the row - update subject cell
        # Detect subject cell index (usually 2, but may vary with merged cells)
        subject_cell = row.cells[2]

        # Add meeting header + new content
        meeting_date = updates.get('meeting_date', '')
        header_text = f"Meeting {meeting_date}" if meeting_date else f"Meeting N{new_meeting_number}"

        add_paragraphs_to_cell(subject_cell, [header_text] + updates['subject_lines'], bold=True)

        # Update for_whom if provided
        if updates.get('for_whom'):
            forwho_cell = row.cells[3]
            # Add empty paragraphs to align, then the new value
            subject_para_count = len(subject_cell.paragraphs)
            forwho_para_count = len(forwho_cell.paragraphs)
            # Pad with empty paragraphs to align
            while forwho_para_count < subject_para_count - 1:
                forwho_cell.add_paragraph('')
                forwho_para_count += 1
            add_paragraphs_to_cell(forwho_cell, [updates['for_whom']], bold=True)

        # Update due if provided
        if updates.get('due'):
            due_cell = row.cells[-1]
            subject_para_count = len(subject_cell.paragraphs)
            due_para_count = len(due_cell.paragraphs)
            while due_para_count < subject_para_count - 1:
                due_cell.add_paragraph('')
                due_para_count += 1
            add_paragraphs_to_cell(due_cell, [updates['due']], bold=True)

        return True

    return False


def _get_tc_text_simple(tc):
    """Get plain text from a table cell XML element."""
    paras = tc.findall(qn('w:p'))
    texts = []
    for p in paras:
        run_texts = [t.text or '' for t in p.findall('.//' + qn('w:t'))]
        texts.append(''.join(run_texts))
    return '\n'.join(texts)


def clone_table_row(table, source_row_idx=-1):
    """Clone a row from a table and append it.

    Copies the full XML structure including cell formatting.
    Returns the new row element.
    """
    source_tr = table.rows[source_row_idx]._tr
    new_tr = copy.deepcopy(source_tr)

    # Clear all text content in the cloned row
    for tc in new_tr.findall(qn('w:tc')):
        for p in tc.findall(qn('w:p')):
            # Remove all runs but keep the paragraph (preserves formatting)
            for r in p.findall(qn('w:r')):
                p.remove(r)

    # Append to table
    table._tbl.append(new_tr)
    return new_tr


def add_new_point(table, point_number, title, subject_lines, for_whom, due,
                  meeting_date='', bold=True):
    """Add a new point row to a section table.

    Clones the last data row for formatting, then fills in content.
    """
    # Clone last row
    new_tr = clone_table_row(table)

    # Get cells from the new row
    tcs = new_tr.findall(qn('w:tc'))
    if len(tcs) < 3:
        return

    # Fill in N° cell
    _set_tc_text(tcs[0], point_number, bold=bold)

    # Fill in Title cell
    _set_tc_text(tcs[1], title, bold=bold)

    # Fill in Subject cell with meeting header + lines
    header = f"Meeting {meeting_date}" if meeting_date else "New point"
    all_lines = [header] + subject_lines
    _set_tc_paragraphs(tcs[2], all_lines, bold=bold)

    # Fill in For whom cell (may be at index 3 or adjusted for gridSpan)
    forwho_idx = len(tcs) - 2
    _set_tc_text(tcs[forwho_idx], for_whom, bold=bold)

    # Fill in Due cell
    due_idx = len(tcs) - 1
    _set_tc_text(tcs[due_idx], due, bold=bold)


def _set_tc_text(tc, text, bold=False):
    """Set text in a table cell XML element, preserving paragraph structure."""
    paras = tc.findall(qn('w:p'))
    if not paras:
        p = OxmlElement('w:p')
        tc.append(p)
        paras = [p]

    # Use first paragraph, add a run
    p = paras[0]
    run = OxmlElement('w:r')
    if bold:
        rPr = OxmlElement('w:rPr')
        b = OxmlElement('w:b')
        rPr.append(b)
        run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    run.append(t)
    p.append(run)


def _set_tc_paragraphs(tc, lines, bold=False):
    """Set multiple paragraphs in a table cell."""
    paras = tc.findall(qn('w:p'))

    for i, text in enumerate(lines):
        if i < len(paras):
            p = paras[i]
        else:
            p = OxmlElement('w:p')
            tc.append(p)

        run = OxmlElement('w:r')
        if bold:
            rPr = OxmlElement('w:rPr')
            b = OxmlElement('w:b')
            rPr.append(b)
            run.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        run.append(t)
        p.append(run)


def find_section_table(doc, section_name):
    """Find the table corresponding to a section name.

    Matching strategy:
    1. Exact section name match in header text
    2. Fallback: any table with subject-type headers (N°, Sujet, Subject)
    """
    # First pass: exact match
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        header_text = ' '.join(c.text for c in table.rows[0].cells)
        if section_name.lower() in header_text.lower():
            return table

    # Second pass: find any subject-type table (for single-table formats like CORUM)
    for table in doc.tables:
        if len(table.rows) < 1:
            continue
        header_text = ' '.join(c.text.lower() for c in table.rows[0].cells)
        is_subject = (
            'subject' in header_text or 'sujet' in header_text
            or 'n°' in header_text or 'n\u00b0' in header_text
            or re.search(r'd\d+\s*[-–]', header_text)
        )
        if is_subject:
            return table

    return None


def _strip_cell_shading(tr):
    """Remove background shading/color from all cells in a row.

    This prevents cloned header rows from inheriting the header color.
    """
    for tc in tr.findall(qn('w:tc')):
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is not None:
            shd = tcPr.find(qn('w:shd'))
            if shd is not None:
                tcPr.remove(shd)


def update_info_exchange(doc, items):
    """Update the information exchange table.

    Args:
        items: list of dicts with {from_whom, status, content, due_date}
    """
    # Find info exchange table
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        header = ' '.join(c.text.lower() for c in table.rows[0].cells)
        if 'from whom' in header or 'de qui' in header:
            # Clear existing data rows
            while len(table.rows) > 1:
                tr = table.rows[-1]._tr
                table._tbl.remove(tr)

            # Add new rows
            for item in items:
                new_tr = clone_table_row(table, 0)  # Clone from header for structure
                _strip_cell_shading(new_tr)
                tcs = new_tr.findall(qn('w:tc'))
                # Clear header text
                for tc in tcs:
                    for p in tc.findall(qn('w:p')):
                        for r in p.findall(qn('w:r')):
                            p.remove(r)

                if len(tcs) >= 4:
                    _set_tc_text(tcs[0], item['from_whom'])
                    _set_tc_text(tcs[1], item['status'])
                    _set_tc_text(tcs[2], item['content'])
                    _set_tc_text(tcs[3], item['due_date'])
            return


def update_planning(doc, items):
    """Update the planning table.

    Args:
        items: list of dicts with {content, is_new}
    """
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        header = table.rows[0].cells[0].text.lower()
        if 'planning' in header and len(table.columns) == 1:
            # Clear existing data rows
            while len(table.rows) > 1:
                tr = table.rows[-1]._tr
                table._tbl.remove(tr)

            # Add new rows
            for item in items:
                new_tr = clone_table_row(table, 0)
                _strip_cell_shading(new_tr)
                tcs = new_tr.findall(qn('w:tc'))
                for tc in tcs:
                    for p in tc.findall(qn('w:p')):
                        for r in p.findall(qn('w:r')):
                            p.remove(r)

                if tcs:
                    bold = item.get('is_new', False)
                    lines = item['content'].split('\n')
                    _set_tc_paragraphs(tcs[0], lines, bold=bold)
            return


def generate_report(previous_path, output_path, updates):
    """Generate a new meeting report from the previous one + updates.

    Args:
        previous_path: Path to previous meeting report .docx
        output_path: Path for the new report .docx
        updates: dict with:
            - meeting_number: int
            - date: str (DD/MM/YYYY)
            - distribution_date: str (DD/MM/YYYY)
            - next_meeting: str (full text for next meeting line)
            - attendance: list of attendance updates (optional)
            - info_exchange: list of {from_whom, status, content, due_date}
            - planning: list of {content, is_new}
            - point_updates: list of {
                section: str,
                number: str,
                subject_lines: list[str],
                for_whom: str,
                due: str,
                meeting_date: str
              }
            - new_points: list of {
                section: str,
                number: str,
                title: str,
                subject_lines: list[str],
                for_whom: str,
                due: str,
                meeting_date: str
              }
    """
    # Step 1: Copy previous report
    doc = copy_report(previous_path, output_path)

    new_num = updates['meeting_number']

    # Step 2: Un-bold all previous "latest" content
    unbold_all_content(doc)

    # Step 3: Update metadata
    update_metadata(
        doc,
        new_num,
        updates['date'],
        updates.get('distribution_date'),
    )

    # Step 4: Update next meeting if provided
    if updates.get('next_meeting'):
        update_next_meeting(
            doc, updates['next_meeting'], updates.get('next_meeting_time')
        )

    # Step 5: Update existing points
    for pu in updates.get('point_updates', []):
        section_table = find_section_table(doc, pu['section'])
        if section_table:
            update_existing_point(
                section_table,
                pu['number'],
                pu,
                new_num,
            )

    # Step 6: Add new points
    for np_data in updates.get('new_points', []):
        section_table = find_section_table(doc, np_data['section'])
        if section_table:
            add_new_point(
                section_table,
                np_data['number'],
                np_data['title'],
                np_data['subject_lines'],
                np_data['for_whom'],
                np_data['due'],
                np_data.get('meeting_date', ''),
            )

    # Step 7: Update info exchange if provided
    if updates.get('info_exchange'):
        update_info_exchange(doc, updates['info_exchange'])

    # Step 8: Update planning if provided
    if updates.get('planning'):
        update_planning(doc, updates['planning'])

    # Step 9: Save
    doc.save(output_path)
    return output_path


def main():
    """CLI entry point - mainly for testing."""
    if len(sys.argv) < 3:
        print("Usage: python report_generator.py <previous_report.docx> <output.docx> [updates.json]")
        print()
        print("Normally called from run.py or Claude Code, not directly.")
        sys.exit(1)

    import json

    previous_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])

    if len(sys.argv) >= 4:
        with open(sys.argv[3], 'r', encoding='utf-8') as f:
            updates = json.load(f)
    else:
        print("No updates file provided. Creating minimal copy with unbolded content.")
        from report_parser import parse_report
        parsed = parse_report(previous_path)
        old_num = parsed['metadata']['meeting_number'] or 0
        updates = {
            'meeting_number': old_num + 1,
            'date': datetime.now().strftime('%d/%m/%Y'),
        }

    generate_report(previous_path, output_path, updates)
    print(f"Report generated: {output_path}")


if __name__ == "__main__":
    main()
