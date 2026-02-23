"""
Meeting Report Parser - Extracts structured data from .docx meeting reports.

Parses meeting reports (Penta and CORUM formats) into JSON with:
- Metadata (meeting number, date, location, distribution, next meeting)
- Attendance matrix
- Info exchange table
- Planning table
- Subject sections with points (number, title, subject paragraphs, for_whom, due)
"""

import json
import re
import sys
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn


def detect_language(doc):
    """Detect report language from keywords in tables and paragraphs."""
    full_text = " ".join(p.text for p in doc.paragraphs) + " "
    full_text += " ".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    full_text_lower = full_text.lower()

    fr_keywords = ["réunion", "objet", "présent", "excusé", "diffusion", "échéance",
                   "planification", "sujets traités", "annexes", "remarques générales"]
    en_keywords = ["meeting", "subject", "present", "excused", "diffusion", "due",
                   "planning", "discussed subjects", "appendices", "general remarks"]

    fr_score = sum(1 for kw in fr_keywords if kw in full_text_lower)
    en_score = sum(1 for kw in en_keywords if kw in full_text_lower)

    return "FR" if fr_score > en_score else "EN"


def parse_metadata_table(table):
    """Parse table 0 - the header/metadata table with merged cells."""
    metadata = {
        "meeting_number": None,
        "date": None,
        "location": None,
        "distribution_date": None,
    }

    for ri, row in enumerate(table.rows):
        row_text = " ".join(cell.text for cell in row.cells)

        # Row 1 typically has meeting number and date
        if ri == 1:
            # Look for meeting number
            num_match = re.search(r'n[°o�]\s*(\d+)', row_text, re.IGNORECASE)
            if not num_match:
                # FR variant: "Compte-rendu de la réunion" with number in a separate cell
                for cell in row.cells:
                    ct = cell.text.strip()
                    if ct.isdigit():
                        num_match = re.match(r'(\d+)', ct)
                        break
            if num_match:
                metadata["meeting_number"] = int(num_match.group(1))

            # Look for date (DD/MM/YYYY) - also handle "Du DD/MM/YYYY"
            date_match = re.search(r'(\d{2}/\d{2}/\d{4})', row_text)
            if date_match:
                metadata["date"] = date_match.group(1)

        # Row 2 typically has location - use first cell only to avoid merged-cell duplication
        if ri == 2:
            first_cell_text = row.cells[0].text.strip()
            loc_match = re.search(r'Location\s*[:\xa0]+\s*(.+)', first_cell_text)
            if not loc_match:
                loc_match = re.search(r'Lieu\s*[:\xa0]+\s*(.+)', first_cell_text)
            if loc_match:
                metadata["location"] = loc_match.group(1).strip()
            elif first_cell_text:
                metadata["location"] = first_cell_text

        # Row 3 typically has distribution date
        if ri == 3:
            dist_match = re.search(r'(?:Distribution|Diffusion)\s+(?:on\s+|le\s+)?(\d{2}/\d{2}/\d{4})', row_text)
            if dist_match:
                metadata["distribution_date"] = dist_match.group(1)

    return metadata


def parse_attendance_table(table):
    """Parse attendance from rows 4-9 of table 0.

    The attendance table is a two-sided matrix (left and right halves)
    with organization names and X marks for Present/Excused/Invited/Diffusion.
    """
    attendees = []

    for ri in range(4, min(len(table.rows), 10)):
        row = table.rows[ri]
        tr = row._tr
        tcs = tr.findall(qn('w:tc'))

        if ri == 4:
            # Header row - skip
            continue

        # Parse both sides (left half = cells 0-4, right half = cells 5-9)
        for side_start in [0, 5]:
            if side_start >= len(tcs):
                break

            name_tc = tcs[side_start]
            name_text = _get_tc_text(name_tc)
            if not name_text.strip():
                continue

            # Parse organization name and people
            lines = [l.strip() for l in name_text.split("\n") if l.strip()]
            if not lines:
                continue

            org_line = lines[0]  # e.g. "MO – Penta Hotel City Centre"
            people = lines[1:] if len(lines) > 1 else []

            # Parse X marks from subsequent cells
            statuses = []
            status_labels = ["Present", "Excused", "Invited", "Diffusion"]
            for offset in range(1, min(5, len(tcs) - side_start)):
                cell_tc = tcs[side_start + offset]
                cell_text = _get_tc_text(cell_tc)
                cell_lines = cell_text.split("\n")
                # Each line corresponds to a person (first line may be empty header)
                for li, line in enumerate(cell_lines):
                    if "X" in line or "x" in line:
                        person_idx = li - 1  # offset by header line
                        if 0 <= person_idx < len(people) and offset - 1 < len(status_labels):
                            label = status_labels[offset - 1]
                            if person_idx < len(statuses):
                                statuses[person_idx].append(label)
                            else:
                                while len(statuses) <= person_idx:
                                    statuses.append([])
                                statuses[person_idx].append(label)

            attendee = {
                "organization": org_line,
                "people": [],
            }
            for pi, person in enumerate(people):
                status = statuses[pi] if pi < len(statuses) else []
                attendee["people"].append({
                    "name": person,
                    "status": status,
                })
            attendees.append(attendee)

    return attendees


def _get_tc_text(tc):
    """Get text from a table cell XML element."""
    paras = tc.findall(qn('w:p'))
    lines = []
    for p in paras:
        texts = [t.text or '' for t in p.findall('.//' + qn('w:t'))]
        lines.append(''.join(texts))
    return '\n'.join(lines)


def parse_info_exchange_table(table):
    """Parse table 1 - information exchange / actions table.

    Columns: From whom | Status | Content | Due date
    """
    items = []
    for ri in range(1, len(table.rows)):
        row = table.rows[ri]
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) >= 4:
            items.append({
                "from_whom": cells[0],
                "status": cells[1],
                "content": cells[2],
                "due_date": cells[3],
            })
    return items


def parse_planning_table(table):
    """Parse table 2 - planning table.

    Single column with multi-line content per row.
    """
    items = []
    for ri in range(1, len(table.rows)):
        cell = table.rows[ri].cells[0]
        text = cell.text.strip()
        if text:
            # Check for bold (latest additions)
            has_bold = any(
                run.bold for p in cell.paragraphs for run in p.runs if run.bold
            )
            items.append({
                "content": text,
                "has_new_content": has_bold,
            })
    return items


def parse_subject_table(table):
    """Parse a subject table (tables 3-6+).

    Detects section name from header, extracts all points with their
    paragraph-level detail including bold tracking.

    Returns: {section_name, table_index, points: [...]}
    """
    # Detect section name and header row
    # Some tables have a merged section-name row (row 0) + header row (row 1)
    # Others have the section name embedded in the header row itself
    section_name = None
    data_start_row = 1  # default: row 0 is header, data starts at row 1

    first_row_text = table.rows[0].cells[0].text.strip()

    # Check if row 0 is a full-width section title (FR variant)
    tr0 = table.rows[0]._tr
    tcs0 = tr0.findall(qn('w:tc'))
    if len(tcs0) == 1 or (len(tcs0) >= 1 and _get_tc_gridspan(tcs0[0]) >= 4):
        # Row 0 is a merged section title, row 1 is the actual header
        match = re.search(r'D\d+\s*[-–]\s*(.+)', first_row_text)
        if match:
            section_name = match.group(1).strip()
        else:
            section_name = first_row_text
        data_start_row = 2  # skip section title + header row
        header_row_idx = 1
    else:
        header_row_idx = 0
        # Detect section name from header row column 2 (Subject column)
        for cell in table.rows[0].cells:
            text = cell.text.strip()
            match = re.search(r'Subject\s*[–\-]\s*(.+)', text)
            if not match:
                match = re.search(r'Objet\s*[–\-]\s*(.+)', text)
            if not match:
                match = re.search(r'D\d+\s*[-–]\s*(.+)', text)
            if not match:
                # CORUM format: "Sujet > Category" or "Sujet > Administratif"
                match = re.search(r'Sujet\s*>\s*(.+)', text)
            if match:
                section_name = match.group(1).strip()
                break

    if section_name is None:
        section_name = "General"

    col_map = _detect_column_map(table, header_row_idx)

    points = []
    for ri in range(data_start_row, len(table.rows)):
        row = table.rows[ri]
        point = _parse_point_row(row, col_map)
        if point:
            points.append(point)

    return {
        "section_name": section_name,
        "points": points,
    }


def _get_tc_gridspan(tc):
    """Get the gridSpan value of a table cell."""
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is not None:
        gs = tcPr.find(qn('w:gridSpan'))
        if gs is not None:
            return int(gs.get(qn('w:val'), '1'))
    return 1


def _detect_column_map(table, header_row_idx=0):
    """Detect which actual XML cells map to which logical columns.

    Returns a dict mapping logical names to cell indices based on
    the header row analysis.
    """
    header_row = table.rows[header_row_idx]
    tr = header_row._tr
    tcs = tr.findall(qn('w:tc'))

    col_map = {
        "number": 0,
        "title": 1,
        "subject": 2,
        "for_whom": None,
        "due": None,
    }

    # Walk through header cells to find For whom and Due
    for ci, tc in enumerate(tcs):
        text = _get_tc_text(tc).strip().lower()
        if 'for whom' in text or 'pour qui' in text:
            if col_map["for_whom"] is None:
                col_map["for_whom"] = ci
        elif text in ('due', 'échéance', 'deadline', 'pour quand'):
            col_map["due"] = ci
        elif text in ('n° du point', 'n°'):
            col_map["number"] = ci
        elif text in ('titres', 'title', 'titre'):
            col_map["title"] = ci
        elif 'sujet' in text or 'subject' in text:
            col_map["subject"] = ci

    # If for_whom wasn't found, use second-to-last; due = last
    if col_map["for_whom"] is None:
        col_map["for_whom"] = len(tcs) - 2
    if col_map["due"] is None:
        col_map["due"] = len(tcs) - 1

    col_map["total_cells"] = len(tcs)
    return col_map


def _parse_point_row(row, col_map):
    """Parse a single data row from a subject table."""
    tr = row._tr
    tcs = tr.findall(qn('w:tc'))

    if len(tcs) < 3:
        return None

    # Number
    number = _get_tc_text(tcs[col_map["number"]]).strip()
    if not number:
        return None

    # Title
    title = _get_tc_text(tcs[col_map["title"]]).strip()

    # Subject - parse paragraph by paragraph with bold detection
    subject_tc = tcs[col_map["subject"]]
    subject_paragraphs = _parse_paragraphs_with_bold(subject_tc)

    # For whom - sparse paragraphs aligned with subject
    for_whom_tc = tcs[col_map["for_whom"]] if col_map["for_whom"] < len(tcs) else None
    for_whom_paragraphs = []
    if for_whom_tc is not None:
        for_whom_paragraphs = _parse_simple_paragraphs(for_whom_tc)

    # Due - sparse paragraphs aligned with subject
    due_tc = tcs[col_map["due"]] if col_map["due"] < len(tcs) else None
    due_paragraphs = []
    if due_tc is not None:
        due_paragraphs = _parse_simple_paragraphs(due_tc)

    return {
        "number": number,
        "title": title,
        "subject_paragraphs": subject_paragraphs,
        "for_whom_paragraphs": for_whom_paragraphs,
        "due_paragraphs": due_paragraphs,
    }


def _parse_paragraphs_with_bold(tc):
    """Parse paragraphs from a cell, tracking bold status per paragraph.

    Returns list of {text, is_bold, has_bold_runs} dicts.
    A paragraph is considered 'bold' if it has any bold runs (= latest meeting content).
    """
    paras = tc.findall(qn('w:p'))
    result = []
    for p in paras:
        runs = p.findall('.//' + qn('w:r'))
        text_parts = []
        bold_parts = []
        for r in runs:
            rPr = r.find(qn('w:rPr'))
            is_bold = False
            if rPr is not None:
                b_elem = rPr.find(qn('w:b'))
                if b_elem is not None:
                    # Check if bold is explicitly set to false
                    val = b_elem.get(qn('w:val'))
                    is_bold = val != '0' and val != 'false'
            t_elem = r.find(qn('w:t'))
            text = t_elem.text if t_elem is not None else ''
            if text:
                text_parts.append(text)
                bold_parts.append(is_bold)

        full_text = ''.join(text_parts)
        has_bold = any(bold_parts)

        result.append({
            "text": full_text,
            "has_bold": has_bold,
        })
    return result


def _parse_simple_paragraphs(tc):
    """Parse paragraphs from a cell as simple text list."""
    paras = tc.findall(qn('w:p'))
    result = []
    for p in paras:
        texts = [t.text or '' for t in p.findall('.//' + qn('w:t'))]
        result.append(''.join(texts))
    return result


def parse_next_meeting(doc):
    """Extract next meeting info from document paragraphs.

    The info may be split across two paragraphs:
    - "Next meeting:" (Header style)
    - "11/02/2026 at 11:00 - On site ..." (Header style)
    Or may appear as a single paragraph.
    """
    found_label = False
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue

        # Check if this is the "Next meeting" label
        if re.search(r'Next meeting|Prochaine r[eé]union', text, re.IGNORECASE):
            found_label = True
            # The date might be in the same paragraph
            date_match = re.search(
                r'(\d{2}/\d{2}/\d{4})\s+(?:at|à)\s+(\d{1,2}[\xa0h:]+\d{2})',
                text
            )
            if date_match:
                return {
                    "date": date_match.group(1),
                    "time": date_match.group(2).replace('\xa0', ''),
                    "full_text": text,
                }
            continue

        # If we just saw the label, this paragraph has the details
        if found_label:
            date_match = re.search(
                r'(\d{2}/\d{2}/\d{4})\s+(?:at|à)\s+(\d{1,2}[\xa0h:]+\d{2})',
                text
            )
            if date_match:
                return {
                    "date": date_match.group(1),
                    "time": date_match.group(2).replace('\xa0', ''),
                    "full_text": text,
                }
            found_label = False

    return None


def classify_tables(doc):
    """Classify each table by its type based on header content.

    Returns list of (table_index, table_type, table) tuples.
    Types: 'metadata', 'info_exchange', 'planning', 'subject'
    """
    classified = []
    for ti, table in enumerate(doc.tables):
        if len(table.rows) == 0:
            continue

        header_text = " ".join(cell.text.strip().lower() for cell in table.rows[0].cells)

        if ti == 0 or 'meeting' in header_text or 'réunion' in header_text:
            if ti == 0:
                classified.append((ti, 'metadata', table))
                continue

        if 'from whom' in header_text or 'de qui' in header_text:
            classified.append((ti, 'info_exchange', table))
        elif 'planning' in header_text and len(table.columns) == 1:
            classified.append((ti, 'planning', table))
        elif 'n°' in header_text or 'n\u00b0' in header_text or 'n�' in header_text:
            classified.append((ti, 'subject', table))
        elif 'sujet' in header_text:
            # CORUM format: table header contains "Sujet"
            classified.append((ti, 'subject', table))
        elif re.search(r'd\d+\s*[-–]', header_text):
            # FR variant: "D1 - Partie Architecture / Sécurité"
            classified.append((ti, 'subject', table))
        elif ti == 0:
            classified.append((ti, 'metadata', table))
        else:
            # Try detecting by structure
            if len(table.columns) == 1 and len(table.rows) > 1:
                classified.append((ti, 'planning', table))
            elif len(table.columns) == 4:
                classified.append((ti, 'info_exchange', table))
            else:
                classified.append((ti, 'unknown', table))

    return classified


def parse_report(docx_path):
    """Parse a meeting report .docx file into structured JSON.

    Args:
        docx_path: Path to the .docx file

    Returns:
        dict with all extracted data
    """
    doc = Document(docx_path)

    language = detect_language(doc)
    next_meeting = parse_next_meeting(doc)

    classified = classify_tables(doc)

    result = {
        "source_file": str(docx_path),
        "language": language,
        "metadata": {},
        "next_meeting": next_meeting,
        "attendance": [],
        "info_exchange": [],
        "planning": [],
        "sections": [],
    }

    for ti, ttype, table in classified:
        if ttype == 'metadata':
            result["metadata"] = parse_metadata_table(table)
            result["attendance"] = parse_attendance_table(table)
        elif ttype == 'info_exchange':
            result["info_exchange"] = parse_info_exchange_table(table)
        elif ttype == 'planning':
            result["planning"] = parse_planning_table(table)
        elif ttype == 'subject':
            section = parse_subject_table(table)
            section["table_index"] = ti
            result["sections"].append(section)

    return result


def main():
    """CLI entry point."""
    if len(sys.argv) < 2:
        print("Usage: python report_parser.py <path_to_report.docx> [output.json]")
        sys.exit(1)

    docx_path = Path(sys.argv[1])
    if not docx_path.exists():
        print(f"Error: File not found: {docx_path}")
        sys.exit(1)

    result = parse_report(docx_path)

    # Output to file or stdout
    if len(sys.argv) >= 3:
        output_path = Path(sys.argv[2])
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, indent=2, ensure_ascii=False)
        print(f"Parsed report saved to: {output_path}")
    else:
        print(json.dumps(result, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
